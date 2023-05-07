import argparse

from dataclasses import dataclass
from datetime import date

from datetime import datetime
from typing import Protocol, Iterable, Sequence
from xml.etree import ElementTree as ET

from docx import Document  # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH  # type: ignore

from pathlib import Path
from urllib import request

_EMPTY_STRING: str = ""
_EMPTY_DATE: date = date(1, 1, 1)


class IData(Protocol):
    first_name: str
    second_name: str
    surname: str
    health_number: str
    ident_number: str
    begin_date: date
    end_date: date


@dataclass
class HealthInfo:
    health_number: str = _EMPTY_STRING
    ident_number: str = _EMPTY_STRING
    first_name: str = _EMPTY_STRING
    second_name: str = _EMPTY_STRING
    surname: str = _EMPTY_STRING
    begin_date: date = _EMPTY_DATE
    end_date: date = _EMPTY_DATE


def load_health_info(path: str | bytes | Path) -> Sequence[IData]:
    _DATE_TIME_FORMAT_STRING: str = "%Y-%m-%dT%H:%M:%S"
    match path:
        case str():
            tree = ET.parse(path)
            root = tree.getroot()
        case Path():
            tree = ET.parse(path)
            root = tree.getroot()
        case bytes():
            root = ET.fromstring(path.decode())
        case _:
            raise ValueError(f"path can be only str or bytes or Path like object")

    health_info: list[HealthInfo] = []
    for e in root.iter('Row'):
        health_number: str = _EMPTY_STRING
        ident_number: str = _EMPTY_STRING
        first_name: str = _EMPTY_STRING
        second_name: str = _EMPTY_STRING
        surname: str = _EMPTY_STRING
        begin_date: date = _EMPTY_DATE
        end_date: date = _EMPTY_DATE
        for child in e:
            if child.tag == 'WIC_NUM':
                health_number = child.text if not child.text is None else _EMPTY_STRING
            if child.tag == 'NP_NUMIDENT':
                ident_number = child.text if not child.text is None else _EMPTY_STRING
            if child.tag == 'NP_SURNAME':
                first_name = child.text if not child.text is None else _EMPTY_STRING
            if child.tag == 'NP_NAME':
                second_name = child.text if not child.text is None else _EMPTY_STRING
            if child.tag == 'NP_PATRONYMIC':
                surname = child.text if not child.text is None else _EMPTY_STRING
            if child.tag == 'WIC_DT_BEGIN':
                begin_date = datetime.strptime(child.text if not child.text is None else _EMPTY_STRING,
                                               _DATE_TIME_FORMAT_STRING)
            if child.tag == 'WIC_DT_END':
                end_date = datetime.strptime(child.text if not child.text is None else _EMPTY_STRING,
                                             _DATE_TIME_FORMAT_STRING)
        health_info.append(
            HealthInfo(health_number, ident_number, first_name, second_name, surname, begin_date, end_date))
    return health_info


def save_health_info(source: str, target: str) -> None:
    document = Document()
    p = document.add_paragraph('РЕЄСТР')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph('листків непрацездатності, які передані зі служби управління персоналом')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p = document.add_paragraph('до комісії зі соціального страхування КП МТК <<Калинівський ринок>>')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    health_info = load_health_info(source)

    table = document.add_table(rows=1, cols=4)
    table_header = table.rows[0].cells
    table_header[0].text = '№ з/п'
    table_header[1].text = 'Прізвище,ім\'я, по батькові'
    table_header[2].text = '№ листка непрацездатності'
    table_header[3].text = 'Номер страхового свідоцтва'

    for index, r in enumerate(health_info):
        row = table.add_row().cells
        row[0].text = f"{index + 1}"
        row[1].text = f"{r.first_name} {r.second_name} {r.surname}"
        row[2].text = f"{r.health_number}"
        row[3].text = f"{r.ident_number}"

    p = document.add_paragraph(f"Листки непрацездатності в кількості {len(health_info)} штуки")
    p = document.add_paragraph(f"Передала:              Прийняв:")
    p = document.add_paragraph(f"Начальник служби       Голова комісії із соціального")
    p = document.add_paragraph(f"управління персоналом  страхування")

    document.save(target)


if __name__ == '__main__':
    parser = argparse.ArgumentParser(
        prog="Xml to docx converter",
        description="Convert xml file to docx as register of sick leave certificates"
    )

    input_group = parser.add_mutually_exclusive_group()

    # Configure parser
    input_group.add_argument(
        "-i", "--input", help="input file in format xml. Default value: export.xml", default='export.xml'
    )

    input_group.add_argument(
        "-u", "--url", help="url address in format http://service.com:<port>/export"
    )

    parser.add_argument(
        "-o", "--output", help="output file. file name for saving data. Default value: reestr.docx",
        default='reestr.docx'
    )

    args = parser.parse_args()
    if args.input:
        save_health_info(args.input, args.output)
    if args.url:
        with request.urlopen(args.url) as reader:
            save_health_info(reader.read().decode('utf-8'), args.output)
